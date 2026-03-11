#!/usr/bin/env python3
"""
OrderChecker v4.0 - Умная система сверки с LLM-парсером

Новые функции v4:
- 🤖 **Ollama интеграция** - локальный LLM для сложных документов
- 🎯 Гибридный парсер: правила → адаптация → LLM
- 💾 Экспорт в Excel, CSV, JSON
- 🔍 Умное определение структуры для всех форматов
- 📊 Расширенная статистика и аналитика
- 🎨 Улучшенный интерфейс с прогрессом

Требования:
    pip install customtkinter pandas openpyxl python-docx pdfplumber requests

Для LLM (опционально):
    # Установить Ollama: https://ollama.ai
    ollama pull llama3.2  # или mistral, qwen2.5
"""

import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
from docx import Document
from pathlib import Path
import threading
import re
import json
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple, Union
from concurrent.futures import ThreadPoolExecutor, as_completed

# PDF поддержка
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# openpyxl для анализа Excel структуры
try:
    import openpyxl
    OPENPYXL_SUPPORT = True
except ImportError:
    OPENPYXL_SUPPORT = False

# Ollama для умного парсинга
try:
    import requests
    REQUESTS_SUPPORT = True
except ImportError:
    REQUESTS_SUPPORT = False

# Настройка
# Настройка
ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")

# Ollama конфигурация
OLLAMA_CONFIG = {
    'base_url': 'http://localhost:11434',
    'model': 'llama3.2',  # или mistral, qwen2.5, neogs
    'timeout': 30,
    'enabled': False  # По умолчанию выключено
}

# ==================== ЛОГИРОВАНИЕ ====================

import logging
from pathlib import Path as PathLib

# Настройка логирования
LOG_DIR = PathLib.home() / ".orderchecker_logs"
LOG_DIR.mkdir(exist_ok=True)

LOG_FILE = LOG_DIR / f"orderchecker_{datetime.now().strftime('%Y%m%d')}.log"

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(LOG_FILE, encoding='utf-8'),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger(__name__)


# ==================== ВАЛИДАЦИЯ ДАННЫХ ====================

class DataValidator:
    """Строгая валидация извлечённых данных"""

    @staticmethod
    def validate_vin(vin: str) -> Tuple[bool, str]:
        """Валидация VIN кода с проверкой контрольной суммы

        Args:
            vin: VIN код для проверки

        Returns:
            (валиден, сообщение_об_ошибке)
        """
        if not vin:
            return False, "VIN пустой"

        vin = vin.upper().strip()

        # Базовая проверка длины и символов
        if len(vin) != 17:
            return False, f"VIN должен быть 17 символов, получено {len(vin)}"

        # Проверка допустимых символов
        invalid_chars = set(vin) - set('0123456789ABCDEFGHJKLMNPRSTUVWXYZ')
        if invalid_chars:
            return False, f"VIN содержит недопустимые символы: {invalid_chars}"

        # Проверка контрольной суммы (для североамериканских VIN)
        # Весовые коэффициенты для позиций
        weights = [8, 7, 6, 5, 4, 3, 2, 10, 0, 9, 8, 7, 6, 5, 4, 3, 2]

        # Значения букв
        transliteration = {
            'A': 1, 'B': 2, 'C': 3, 'D': 4, 'E': 5, 'F': 6, 'G': 7, 'H': 8,
            'J': 1, 'K': 2, 'L': 3, 'M': 4, 'N': 5, 'P': 7, 'R': 9,
            'S': 2, 'T': 3, 'U': 4, 'V': 5, 'W': 6, 'X': 7, 'Y': 8, 'Z': 9
        }

        try:
            total = 0
            for i, char in enumerate(vin):
                if char.isdigit():
                    value = int(char)
                else:
                    value = transliteration.get(char, 0)

                if i < 9:  # Пропускаем 9-ю позицию (контрольную цифру)
                    total += value * weights[i]

            check_digit = vin[8]
            if check_digit == 'X':
                expected = 10
            else:
                expected = total % 11
                if expected == 10:
                    expected = 'X'
                else:
                    expected = str(expected)

            if check_digit != expected:
                logger.warning(f"VIN контрольная сумма не совпадает: ожидается {expected}, получено {check_digit}")
                # Не считаем ошибкой, но предупреждаем

        except Exception as e:
            logger.warning(f"Ошибка проверки контрольной суммы VIN: {e}")

        return True, "VIN корректен"

    @staticmethod
    def validate_plate(plate: str) -> Tuple[bool, str]:
        """Валидация российского гос. номера

        Args:
            plate: гос. номер для проверки

        Returns:
            (валиден, сообщение_об_ошибке)
        """
        if not plate:
            return False, "Гос. номер пустой"

        plate = plate.upper().strip()

        # Российский формат: А000АА00 или А000АА000
        pattern1 = r'^[АВЕКМНОРСТУХ]\d{3}[АВЕКМНОРСТУХ]{2}\d{2,3}$'

        if not re.match(pattern1, plate):
            return False, f"Гос. номер не соответствует формату: {plate}"

        return True, "Гос. номер корректен"

    @staticmethod
    def validate_order_number(order_num: str) -> Tuple[bool, str]:
        """Валидация номера заказа

        Args:
            order_num: номер заказа для проверки

        Returns:
            (валиден, сообщение_об_ошибке)
        """
        if not order_num:
            return False, "Номер заказа пустой"

        order_num = str(order_num).strip()

        # Проверяем что это 8 цифр
        if not re.match(r'^\d{8}$', order_num):
            return False, f"Номер заказа должен быть 8 цифр, получено: {order_num}"

        # Дополнительная проверка: первая цифра обычно 0-9, но не 0
        if order_num[0] == '0':
            logger.warning(f"Номер заказа начинается с 0: {order_num}")

        return True, "Номер заказа корректен"

    @staticmethod
    def validate_date(date_str: str) -> Tuple[bool, str, Optional[datetime]]:
        """Валидация даты

        Args:
            date_str: строка с датой

        Returns:
            (валиден, сообщение, datetime_объект)
        """
        if not date_str:
            return False, "Дата пустая", None

        date_formats = [
            '%d.%m.%Y',
            '%Y-%m-%d',
            '%d/%m/%Y',
            '%d.%m.%y',
            '%Y/%m/%d'
        ]

        for fmt in date_formats:
            try:
                dt = datetime.strptime(date_str.strip(), fmt)

                # Проверка разумности даты
                if dt.year < 2000 or dt.year > datetime.now().year + 1:
                    return False, f"Дата вне разумного диапазона: {dt.year}", dt

                return True, "Дата корректна", dt
            except ValueError:
                continue

        return False, f"Не удалось распознать дату: {date_str}", None

    @staticmethod
    def validate_amount(amount_str: str) -> Tuple[bool, str, Optional[float]]:
        """Валидация суммы

        Args:
            amount_str: строка с суммой

        Returns:
            (валиден, сообщение, float_сумма)
        """
        if not amount_str:
            return False, "Сумма пустая", None

        # Удаляем все кроме цифр и точки
        cleaned = re.sub(r'[^\d.]', '', str(amount_str))

        if not cleaned:
            return False, f"Не удалось извлечь сумму из: {amount_str}", None

        try:
            amount = float(cleaned)

            # Проверка разумности суммы
            if amount < 0:
                return False, "Сумма не может быть отрицательной", amount
            if amount > 10_000_000:
                logger.warning(f"Очень большая сумма: {amount}")

            return True, "Сумма корректна", amount
        except ValueError:
            return False, f"Не удалось преобразовать в число: {cleaned}", None

    @staticmethod
    def validate_order(order: Dict[str, Any]) -> Dict[str, Any]:
        """Полная валидация заказа

        Args:
            order: словарь с данными заказа

        Returns:
            словарь с результатами валидации
        """
        validation_result = {
            'valid': True,
            'errors': [],
            'warnings': [],
            'fields': {}
        }

        # Валидация номера заказа (обязательное поле)
        if order.get('order_number'):
            is_valid, msg = DataValidator.validate_order_number(order['order_number'])
            validation_result['fields']['order_number'] = {'valid': is_valid, 'message': msg}
            if not is_valid:
                validation_result['valid'] = False
                validation_result['errors'].append(f"Номер заказа: {msg}")
        else:
            validation_result['valid'] = False
            validation_result['errors'].append("Отсутствует номер заказа")

        # Валидация VIN
        if order.get('vin'):
            is_valid, msg = DataValidator.validate_vin(order['vin'])
            validation_result['fields']['vin'] = {'valid': is_valid, 'message': msg}
            if not is_valid:
                validation_result['errors'].append(f"VIN: {msg}")

        # Валидация гос. номера
        if order.get('plate'):
            is_valid, msg = DataValidator.validate_plate(order['plate'])
            validation_result['fields']['plate'] = {'valid': is_valid, 'message': msg}
            if not is_valid:
                validation_result['errors'].append(f"Гос. номер: {msg}")

        # Валидация даты
        if order.get('date_open'):
            is_valid, msg, dt = DataValidator.validate_date(order['date_open'])
            validation_result['fields']['date_open'] = {'valid': is_valid, 'message': msg}
            if not is_valid:
                validation_result['errors'].append(f"Дата: {msg}")

        # Валидация суммы
        if order.get('amount'):
            is_valid, msg, amount = DataValidator.validate_amount(order['amount'])
            validation_result['fields']['amount'] = {'valid': is_valid, 'message': msg}
            if not is_valid:
                validation_result['errors'].append(f"Сумма: {msg}")

        return validation_result


# Цвета
COLORS = {
    'primary': '#10B981',  # Зелёный
    'success': '#10B981',  # Зелёный
    'warning': '#F59E0B',  # Оранжевый
    'error': '#EF4444',    # Красный
    'info': '#F59E0B',     # Оранжевый
}


# ==================== БАЗОВЫЙ ПАРСЕР ====================

class BaseParser:
    """Базовый класс для парсеров документов"""

    PATTERNS = {
        'order_number': [
            r'Заказ-наряд\s*№\s*(\d{8})',
            r'РЕМОНТНЫЙ\s*ЗАКАЗ\s*№\s*(\d{8})',
            r'Счет-Заказ\s*№\s*(\d{8})',
            r'заказ\s*[№N]?\s*(\d{8})',
        ],
        'vin': [
            r'VIN[:\s]*([A-HJ-NPR-Z0-9]{17})',
            r'VIN-код[:\s]*([A-HJ-NPR-Z0-9]{17})',
        ],
        'plate': [
            r'Гос\.?\s*номер[:\s]*([АВЕКМНОРСТУХ]\d{3}[АВЕКМНОРСТУХ]{2}\d{2,3})',
            r'Регистрационный\s*номер[:\s]*([АВЕКМНОРСТУХ]\d{3}[АВЕКМНОРСТУХ]{2}\d{2,3})',
            r'[АВЕКМНОРСТУХ]\d{3}[АВЕКМНОРСТУХ]{2}\d{2,3}',
        ],
        'date': [
            r'(\d{2})\.(\d{2})\.(\d{4})',
        ],
        'amount': [
            r'Итого[:\s]*[\d\s]+руб',
            r'Всего[:\s]*[\d\s]+руб',
            r'(\d[\d\s]*\d)\s*руб',
            r'(\d[\d\s]*\d)\s*₽',
        ],
    }

    @staticmethod
    def extract_pattern(text: str, pattern_group: str) -> Optional[str]:
        """Извлекает данные по паттерну"""
        for pattern in BaseParser.PATTERNS.get(pattern_group, []):
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1) if match.lastindex else match.group(0)
        return None

    @staticmethod
    def find_patterns_all(text: str, pattern_group: str) -> List[str]:
        """Находит все совпадения по паттерну"""
        results = []
        for pattern in BaseParser.PATTERNS.get(pattern_group, []):
            results.extend(re.findall(pattern, text, re.IGNORECASE))
        return results

    @staticmethod
    def is_valid_vin(vin: str) -> bool:
        """Проверяет валидность VIN"""
        if not vin or len(vin) != 17:
            return False
        return not any(c in vin for c in 'IOQ')


# ==================== УМНЫЙ ПАРСЕР С LLM ====================

class OllamaClient:
    """Клиент для работы с Ollama API"""

    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or OLLAMA_CONFIG
        self.base_url = self.config.get('base_url', 'http://localhost:11434')
        self.model = self.config.get('model', 'llama3.2')
        self.timeout = self.config.get('timeout', 30)
        self._available = None

    def is_available(self) -> bool:
        """Проверяет доступность Ollama"""
        if self._available is not None:
            return self._available

        if not REQUESTS_SUPPORT:
            self._available = False
            return False

        try:
            response = requests.get(f"{self.base_url}/api/tags", timeout=5)
            self._available = response.status_code == 200
            return self._available
        except Exception:
            self._available = False
            return False

    def extract_order_data(self, text: str, filename: str = "") -> Dict[str, Any]:
        """Извлекает данные заказа с помощью LLM"""

        prompt = f"""Из следующего текста документа извлеки данные о заказе на ремонт автомобиля.

Правила:
1. Верни ТОЛЬКО валидный JSON
2. Все поля обязательны, если данных нет - верни пустую строку ""
3. VIN - это 17 символов (латинские буквы и цифры, без I, O, Q)
4. Гос. номер - формат: А000АА00 или А000АА000 (русские буквы)
5. Дата - формат: ДД.ММ.ГГГГ
6. Сумму пиши без пробелов, только цифры

Формат ответа (строго JSON):
{{
    "order_number": "8-значный номер заказа",
    "date_open": "дата открытия или пусто",
    "date_close": "дата закрытия или пусто",
    "vin": "VIN код или пусто",
    "plate": "гос. номер или пусто",
    "model": "модель авто или пусто",
    "amount": "сумма рублями или пусто"
}}

Текст документа:
{text[:8000]}  Ограничение: 8000 символов

Ответ:"""

        try:
            response = requests.post(
                f"{self.base_url}/api/generate",
                json={
                    "model": self.model,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.1,  # Минимальная креативность
                        "num_predict": 500,
                    }
                },
                timeout=self.timeout
            )

            if response.status_code == 200:
                result = response.json()
                llm_text = result.get('response', '')

                # Извлекаем JSON из ответа
                json_match = re.search(r'\{[^{}]+\}', llm_text, re.DOTALL)
                if json_match:
                    data = json.loads(json_match.group())

                    # Валидация
                    if data.get('order_number') and len(str(data['order_number'])) >= 6:
                        # Форматируем сумму
                        if data.get('amount'):
                            amount = re.sub(r'[^\d]', '', str(data['amount']))
                            if amount:
                                data['amount'] = f"{amount} руб."

                        return {
                            'order_number': str(data['order_number']),
                            'date_open': data.get('date_open', ''),
                            'date_close': data.get('date_close', ''),
                            'vin': data.get('vin', ''),
                            'plate': data.get('plate', ''),
                            'model': data.get('model', ''),
                            'amount': data.get('amount', ''),
                            'source_info': {
                                'order_number': '🤖 LLM',
                                'date_open': '🤖 LLM',
                                'date_close': '🤖 LLM',
                                'vin': '🤖 LLM',
                                'plate': '🤖 LLM',
                                'model': '🤖 LLM',
                                'amount': '🤖 LLM'
                            }
                        }

        except Exception as e:
            print(f"Ollama ошибка: {e}")

        return None


class SmartParser:
    """Гибридный умный парсер: правила → адаптация → LLM"""

    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or {}
        self.llm = OllamaClient(config) if self.config.get('llm_enabled', False) else None

    def parse(self, filepath: str) -> List[Dict[str, Any]]:
        """Умный парсинг с многослойным подходом"""

        path = Path(filepath)
        ext = path.suffix.lower()

        # === СЛОЙ 1: Базовые парсеры ===
        orders = self._parse_with_rules(filepath, ext)

        # Если нашли данные - возвращаем
        if orders and self._is_valid_result(orders):
            return orders

        # === СЛОЙ 2: Адаптивный парсинг ===
        if not orders or not self._is_valid_result(orders):
            orders = self._parse_adaptive(filepath, ext)
            if orders and self._is_valid_result(orders):
                return orders

        # === СЛОЙ 3: LLM (если включен и доступен) ===
        if self.llm and self.llm.is_available():
            llm_result = self._parse_with_llm(filepath, ext)
            if llm_result:
                return [llm_result]

        # Возвращаем что нашли (может быть пусто)
        return orders or []

    def _parse_with_rules(self, filepath: str, ext: str) -> List[Dict[str, Any]]:
        """Базовый парсинг по правилам"""
        try:
            if ext == '.docx':
                return DOCXParser(filepath).parse()
            elif ext == '.pdf':
                return PDFParser(filepath).parse()
            elif ext in ['.xlsx', '.xls']:
                return ExcelParser(filepath).parse()
            elif ext == '.csv':
                return CSVParser(filepath).parse()
        except Exception as e:
            print(f"Ошибка базового парсера {ext}: {e}")
        return []

    def _parse_adaptive(self, filepath: str, ext: str) -> List[Dict[str, Any]]:
        """Адаптивный парсинг для сложных случаев"""
        # Для PDF - пробуем разные стратегии извлечения
        if ext == '.pdf' and PDF_SUPPORT:
            import pdfplumber
            try:
                with pdfplumber.open(filepath) as pdf:
                    all_text = ""
                    all_tables = []

                    for page in pdf.pages:
                        # Извлекаем текст
                        text = page.extract_text() or ""
                        all_text += text + "\n"

                        # Извлекаем таблицы
                        tables = page.extract_tables()
                        if tables:
                            all_tables.extend(tables)

                    # Комбинируем данные
                    combined_text = all_text
                    for table in all_tables:
                        for row in table:
                            for cell in row:
                                if cell:
                                    combined_text += str(cell) + " "

                    # Ищем заказы
                    return PDFParser(filepath).parse()
            except Exception as e:
                print(f"Адаптивный PDF парсинг: {e}")

        return []

    def _parse_with_llm(self, filepath: str, ext: str) -> Optional[Dict[str, Any]]:
        """Парсинг с помощью LLM"""
        if not self.llm:
            return None

        # Извлекаем текст из файла
        text = self._extract_text_from_file(filepath, ext)
        if not text:
            return None

        # Отправляем в LLM
        return self.llm.extract_order_data(text, Path(filepath).name)

    def _extract_text_from_file(self, filepath: str, ext: str) -> Optional[str]:
        """Извлекает текст из файла для LLM"""
        try:
            if ext == '.docx':
                doc = Document(filepath)
                text = '\n'.join([p.text for p in doc.paragraphs])
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                return text

            elif ext == '.pdf' and PDF_SUPPORT:
                import pdfplumber
                with pdfplumber.open(filepath) as pdf:
                    text = ""
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        text += page_text + "\n"
                        # Таблицы
                        tables = page.extract_tables()
                        if tables:
                            for table in tables:
                                for row in table:
                                    for cell in row:
                                        if cell:
                                            text += str(cell) + " "
                    return text

            elif ext in ['.xlsx', '.xls']:
                df = pd.read_excel(filepath, dtype=str)
                return df.to_string()

            elif ext == '.csv':
                df = pd.read_csv(filepath, dtype=str)
                return df.to_string()

        except Exception as e:
            print(f"Ошибка извлечения текста: {e}")

        return None

    def _is_valid_result(self, orders: List[Dict[str, Any]]) -> bool:
        """Проверяет валидность результата"""
        if not orders:
            return False

        # Проверяем что есть хотя бы один заказ с номером
        for order in orders:
            if order.get('order_number') and len(str(order['order_number'])) >= 6:
                return True

        return False


# ==================== ПАРСЕРЫ ФАЙЛОВ ====================

class DOCXParser(BaseParser):
    """Парсер DOCX с улучшенными паттернами"""

    def __init__(self, filepath: str):
        self.filepath = Path(filepath)
        self.doc = Document(filepath)

    def parse(self) -> List[Dict[str, Any]]:
        orders = []
        current_order = None
        i = 0

        while i < len(self.doc.tables):
            table = self.doc.tables[i]
            order_number = self._find_order_number(table)

            if order_number:
                if current_order:
                    orders.append(current_order)

                current_order = {
                    'order_number': order_number,
                    'date_open': self._find_date_in_table(table, ['Дата открытия']),
                    'date_close': self._find_date_in_table(table, ['Дата закрытия']),
                    'vin': None,
                    'plate': None,
                    'model': None,
                    'amount': None,
                    # Отслеживание источника данных
                    'source_info': {
                        'order_number': f'Таблица {i + 1}',
                        'date_open': f'Таблица {i + 1}',
                        'date_close': f'Таблица {i + 1}',
                        'plate': f'Таблица {i + 1}' if self._find_plate_in_table(table) else '',
                        'vin': f'Таблица {i + 2}' if i + 1 < len(self.doc.tables) else '',
                        'model': f'Таблица {i + 2}' if i + 1 < len(self.doc.tables) else '',
                        'amount': f'Таблица {i + 2}' if i + 1 < len(self.doc.tables) else ''
                    }
                }

                plate_match = self._find_plate_in_table(table)
                if plate_match:
                    current_order['plate'] = plate_match

                if i + 1 < len(self.doc.tables):
                    next_table = self.doc.tables[i + 1]
                    current_order['vin'] = self._find_vin_in_table(next_table)
                    if not current_order['plate']:
                        current_order['plate'] = self._find_plate_in_table(next_table)
                        current_order['source_info']['plate'] = f'Таблица {i + 2}'
                    current_order['model'] = self._find_model_in_table(next_table)
                    current_order['amount'] = self._find_amount_in_table(next_table)

            i += 1

        if current_order:
            orders.append(current_order)

        return orders

    def _find_order_number(self, table):
        """Поиск номера заказа в таблице"""
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for pattern in self.PATTERNS['order_number']:
                    match = re.search(pattern, text, re.IGNORECASE)
                    if match:
                        return match.group(1)
        return None

    def _find_date_in_table(self, table, keywords):
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells]
            for idx, text in enumerate(cells_text):
                for keyword in keywords:
                    if keyword in text and idx + 1 < len(cells_text):
                        match = re.search(r'(\d{2}\.\d{2}\.\d{4})', cells_text[idx + 1])
                        if match:
                            return match.group(1)
        return None

    def _find_vin_in_table(self, table):
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text == 'VIN' and idx + 1 < len(row.cells):
                    vin = row.cells[idx + 1].text.strip()
                    if self.is_valid_vin(vin):
                        return vin
        return None

    def _find_plate_in_table(self, table):
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for pattern in self.PATTERNS['plate']:
                    match = re.search(pattern, text, re.IGNORECASE)
                    if match:
                        return match.group(1) if match.lastindex else match.group(0)
        return None

    def _find_model_in_table(self, table):
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text == 'Модель' and idx + 1 < len(row.cells):
                    model = row.cells[idx + 1].text.strip()
                    if model:
                        return model
        return None

    def _find_amount_in_table(self, table):
        for row in table.rows:
            row_text = ' '.join([cell.text for cell in row.cells])
            for pattern in self.PATTERNS['amount']:
                match = re.search(pattern, row_text, re.IGNORECASE)
                if match:
                    amount = re.sub(r'[^\d]', '', match.group(1) if match.lastindex else match.group(0))
                    try:
                        return f"{int(amount)} руб."
                    except ValueError:
                        continue
        return None


class PDFParser(BaseParser):
    """Улучшенный парсер PDF с поддержкой таблиц"""

    def __init__(self, filepath: str):
        if not PDF_SUPPORT:
            raise ImportError("pdfplumber не установлен")
        self.filepath = Path(filepath)

    def parse(self) -> List[Dict[str, Any]]:
        orders = []

        with pdfplumber.open(self.filepath) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                # Извлекаем таблицы если есть
                tables = page.extract_tables()
                tables_text = ""
                if tables:
                    for table in tables:
                        for row in table:
                            for cell in row:
                                if cell:
                                    tables_text += str(cell) + " "

                # Извлекаем текст
                text = page.extract_text() or ""

                # Комбинируем
                full_text = text + " " + tables_text

                # Ищем заказы улучшенными методами
                orders.extend(self._extract_orders_from_text(full_text, page_num))

        # Убираем дубликаты
        seen = set()
        unique_orders = []
        for order in orders:
            if order['order_number'] not in seen:
                seen.add(order['order_number'])
                unique_orders.append(order)

        return unique_orders

    def _extract_orders_from_text(self, text: str, page_num: int = 1) -> List[Dict[str, Any]]:
        orders = []
        order_numbers = []

        # Строгие паттерны
        for pattern in self.PATTERNS['order_number']:
            order_numbers.extend(re.findall(pattern, text, re.IGNORECASE))

        # Если не нашли строгие паттерны - ищем 8-значные числа
        if not order_numbers:
            lines = text.split('\n')
            for line in lines:
                if any(kw in line.upper() for kw in ['ЗАКАЗ', 'РЕМОНТ', 'НАРЯД']):
                    match = re.search(r'\b(\d{8})\b', line)
                    if match and not (match.group(1).startswith('20') or match.group(1).startswith('19')):
                        order_numbers.append(match.group(1))

        # Убираем дубликаты
        seen = set()
        for num in order_numbers:
            if num not in seen:
                seen.add(num)
                orders.append({
                    'order_number': num,
                    'date_open': self.extract_pattern(text, 'date'),
                    'date_close': None,
                    'vin': None,
                    'plate': self.extract_pattern(text, 'plate'),
                    'model': None,
                    'amount': self.extract_pattern(text, 'amount'),
                    'source_info': {
                        'order_number': f'Страница {page_num}',
                        'date_open': f'Страница {page_num}',
                        'plate': f'Страница {page_num}',
                        'amount': f'Страница {page_num}'
                    }
                })

        return orders


class ExcelParser(BaseParser):
    """Парсер Excel файлов (как источники заказов)"""

    def __init__(self, filepath: str):
        self.filepath = Path(filepath)

    def parse(self) -> List[Dict[str, Any]]:
        orders = []

        try:
            # Пытаемся определить структуру
            df = pd.read_excel(self.filepath, sheet_name=0, header=None, dtype=str)

            # Ищем строку заголовка по ключевым словам
            header_row = 0
            for idx, row in df.iterrows():
                row_text = ' '.join([str(v) for v in row.values]).lower()
                if any(kw in row_text for kw in ['номер', 'заказ', 'vin', 'гос']):
                    header_row = idx
                    break

            # Перечитываем с правильным заголовком
            df = pd.read_excel(self.filepath, sheet_name=0, header=header_row, dtype=str)

            # Ищем колонку с номерами заказов
            order_col = None
            for col in df.columns:
                col_str = str(col).lower()
                if 'номер' in col_str or 'заказ' in col_str:
                    order_col = col
                    break

            # Если не нашли по названию, берем вторую колонку
            if order_col is None and len(df.columns) > 1:
                order_col = df.columns[1]

            if order_col is not None:
                for idx, row in df.iterrows():
                    excel_row = idx + header_row + 2  # +1 для 0-based, +1 для заголовка
                    order_num = str(row.get(order_col, ''))
                    order_num_match = re.search(r'(\d{8})', order_num)
                    if order_num_match:
                        orders.append({
                            'order_number': order_num_match.group(1),
                            'date_open': str(row.get('Дата', row.get('дата', '')))[:10],
                            'date_close': None,
                            'vin': str(row.get('VIN', '')),
                            'plate': str(row.get('Гос', row.get('Гос. номер', row.get('гос. номер', '')))),
                            'model': str(row.get('Модель', row.get('модель', ''))),
                            'amount': str(row.get('Сумма', row.get('сумма', ''))),
                            'source_info': {
                                'order_number': f'Строка {excel_row}',
                                'date_open': f'Строка {excel_row}',
                                'vin': f'Строка {excel_row}',
                                'plate': f'Строка {excel_row}',
                                'model': f'Строка {excel_row}',
                                'amount': f'Строка {excel_row}'
                            }
                        })

        except Exception as e:
            print(f"Ошибка чтения Excel: {e}")

        return orders


class CSVParser(BaseParser):
    """Парсер CSV файлов"""

    def __init__(self, filepath: str):
        self.filepath = Path(filepath)

    def parse(self) -> List[Dict[str, Any]]:
        orders = []

        try:
            # Пробуем разные кодировки
            for encoding in ['utf-8-sig', 'utf-8', 'cp1251']:
                try:
                    df = pd.read_csv(self.filepath, encoding=encoding, dtype=str)
                    break
                except:
                    continue
            else:
                df = pd.read_csv(self.filepath, encoding='utf-8', dtype=str, errors='ignore')

            # Ищем колонку с номерами заказов
            order_col = None
            for col in df.columns:
                col_str = str(col).lower()
                if 'номер' in col_str or 'заказ' in col_str:
                    order_col = col
                    break

            if order_col is None and len(df.columns) > 1:
                order_col = df.columns[1]

            if order_col is not None:
                for idx, row in df.iterrows():
                    csv_row = idx + 2  # +1 для 0-based, +1 для заголовка
                    order_num = str(row.get(order_col, ''))
                    order_num_match = re.search(r'(\d{8})', order_num)
                    if order_num_match:
                        orders.append({
                            'order_number': order_num_match.group(1),
                            'date_open': str(row.get('Дата', row.get('дата', '')))[:10],
                            'date_close': None,
                            'vin': str(row.get('VIN', '')),
                            'plate': str(row.get('Гос', row.get('Гос. номер', ''))),
                            'model': str(row.get('Модель', row.get('модель', ''))),
                            'amount': str(row.get('Сумма', row.get('сумма', ''))),
                            'source_info': {
                                'order_number': f'Строка {csv_row}',
                                'date_open': f'Строка {csv_row}',
                                'vin': f'Строка {csv_row}',
                                'plate': f'Строка {csv_row}',
                                'model': f'Строка {csv_row}',
                                'amount': f'Строка {csv_row}'
                            }
                        })
        except Exception as e:
            print(f"Ошибка чтения CSV: {e}")

        return orders


class UniversalParser:
    """Универсальный парсер для всех форматов документов"""

    @staticmethod
    def parse(filepath: str) -> List[Dict[str, Any]]:
        """Универсальный парсинг - определяет формат автоматически"""
        path = Path(filepath)
        ext = path.suffix.lower()

        if ext == '.docx':
            return DOCXParser(filepath).parse()
        elif ext == '.pdf':
            return PDFParser(filepath).parse()
        elif ext in ['.xlsx', '.xls']:
            return ExcelParser(filepath).parse()
        elif ext == '.csv':
            return CSVParser(filepath).parse()
        else:
            return []


# ==================== ОБУЧЕНИЕ НА РЕЕСТРЕ ====================

class PatternStorage:
    """Хранилище паттернов с сохранением на диск"""

    def __init__(self, storage_path: str = None):
        """Инициализация хранилища

        Args:
            storage_path: путь к файлу для сохранения паттернов (по умолчанию в домашней директории)
        """
        if storage_path is None:
            home = Path.home()
            storage_path = home / ".orderchecker_patterns.json"

        self.storage_path = Path(storage_path)
        self.patterns = {}
        self.metadata = {}

        # Загружаем при создании
        self._load()

    def save(self, patterns: Dict[str, List[str]], metadata: Dict[str, Any] = None):
        """Сохраняет паттерны в файл

        Args:
            patterns: словарь {тип: [список паттернов]}
            metadata: дополнительная информация (дата реестра, количество записей и т.д.)
        """
        self.patterns = patterns
        self.metadata = metadata or {}

        try:
            data = {
                'patterns': self.patterns,
                'metadata': self.metadata,
                'saved_at': datetime.now().isoformat(),
                'version': '1.0'
            }

            with open(self.storage_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

            print(f"💾 Паттерны сохранены: {self.storage_path}")
            print(f"   Всего паттернов: {sum(len(p) for p in self.patterns.values())}")

        except Exception as e:
            print(f"❌ Ошибка сохранения паттернов: {e}")

    def _load(self):
        """Загружает паттерны из файла"""
        if not self.storage_path.exists():
            print(f"📁 Хранилище не найдено: {self.storage_path}")
            return

        try:
            with open(self.storage_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            self.patterns = data.get('patterns', {})
            self.metadata = data.get('metadata', {})

            print(f"📂 Паттерны загружены: {self.storage_path}")
            print(f"   Всего паттернов: {sum(len(p) for p in self.patterns.values())}")
            if self.metadata:
                print(f"   Метаданные: {self.metadata}")

        except Exception as e:
            print(f"❌ Ошибка загрузки паттернов: {e}")
            self.patterns = {}
            self.metadata = {}

    def get_patterns(self) -> Dict[str, List[str]]:
        """Возвращает загруженные паттерны"""
        return self.patterns

    def get_flat_patterns(self) -> List[str]:
        """Возвращает плоский список всех паттернов"""
        all_patterns = []
        for patterns in self.patterns.values():
            all_patterns.extend(patterns)
        return list(set(all_patterns))

    def is_empty(self) -> bool:
        """Проверяет есть ли сохранённые паттерны"""
        return len(self.patterns) == 0

    def needs_update(self, registry_path: str, record_count: int = None) -> bool:
        """Проверяет нужно ли обновить паттерны

        Args:
            registry_path: путь к текущему реестру
            record_count: количество записей в реестре

        Returns:
            True если нужно обновить, False если паттерны актуальны
        """
        if self.is_empty():
            logger.info("📭 Паттернов нет, нужно извлечь")
            return True

        # Проверяем путь к реестру
        stored_path = self.metadata.get('registry_path', '')
        if stored_path != str(registry_path):
            logger.info(f"🔄 Реестр изменился: {stored_path} -> {registry_path}")
            return True

        # Проверяем количество записей
        if record_count is not None:
            stored_count = self.metadata.get('record_count', 0)
            if stored_count != record_count:
                logger.info(f"📊 Количество записей изменилось: {stored_count} -> {record_count}")
                return True

        # Проверяем возраст паттернов (старше 7 дней = обновить)
        if 'saved_at' in self.metadata:
            try:
                saved_at = datetime.fromisoformat(self.metadata['saved_at'])
                age_days = (datetime.now() - saved_at).days

                if age_days > 7:
                    logger.info(f"⏰ Паттерны устарели ({age_days} дней)")
                    return True
            except:
                pass

        logger.info("✅ Паттерны актуальны, обновление не требуется")
        return False

    def get_metadata(self) -> Dict[str, Any]:
        """Возвращает метаданные о сохранённых паттернах"""
        return self.metadata.copy()

    def clear(self):
        """Очищает хранилище"""
        self.patterns = {}
        self.metadata = {}
        if self.storage_path.exists():
            self.storage_path.unlink()
            print(f"🗑️ Хранилище очищено: {self.storage_path}")
            logger.info(f"Хранилище очищено: {self.storage_path}")


class PatternExtractor:
    """Извлекает паттерны из реестра для умного поиска в документах"""

    def __init__(self, registry_df: pd.DataFrame):
        """Инициализация с данными реестра"""
        self.df = registry_df
        self.patterns = {}
        self._extract_all_patterns()

    def _extract_all_patterns(self):
        """Извлекает уникальные паттерны из всех колонок реестра"""
        print("🔍 Извлекаю паттерны из реестра...")

        # Проходим по всем колонкам и извлекаем уникальные значения
        for col in self.df.columns:
            col_data = self.df[col].dropna()

            if col_data.empty:
                continue

            # Определяем тип данных по содержимому
            patterns = self._extract_patterns_from_column(col, col_data)

            if patterns:
                self.patterns[col] = patterns
                print(f"  ✓ Колонка '{col}': {len(patterns)} уникальных паттернов")

        print(f"✅ Всего извлечено: {sum(len(p) for p in self.patterns.values())} паттернов из {len(self.patterns)} колонок")

    def _extract_patterns_from_column(self, col_name: str, col_data: pd.Series) -> List[str]:
        """Извлекает паттерны из колонки в зависимости от типа данных"""
        patterns = []

        # Пробуем разные паттерны
        for value in col_data:
            value_str = str(value).strip()

            if not value_str or value_str.lower() in ['nan', 'none', '']:
                continue

            # VIN код (17 символов)
            if self._is_vin(value_str):
                patterns.append(value_str.upper())

            # Гос. номер (формат А000АА00)
            elif self._is_plate(value_str):
                patterns.append(value_str.upper())

            # Номер заказа (8 цифр)
            elif self._is_order_number(value_str):
                patterns.append(value_str)

            # Дата
            elif self._is_date(value_str):
                patterns.append(value_str)

            # Сумма
            elif self._is_amount(value_str):
                patterns.append(value_str)

            # Текстовые данные (модели, услуги и т.д.)
            elif len(value_str) > 2 and len(value_str) < 100:
                patterns.append(value_str)

        return list(set(patterns))  # Уникальные значения

    @staticmethod
    def _is_vin(value: str) -> bool:
        """Проверяет VIN код"""
        vin_pattern = r'^[A-HJ-NPR-Z0-9]{17}$'
        return bool(re.match(vin_pattern, value.upper()))

    @staticmethod
    def _is_plate(value: str) -> bool:
        """Проверяет гос. номер"""
        plate_pattern = r'^[АВЕКМНОРСТУХ]\d{3}[АВЕКМНОРСТУХ]{2}\d{2,3}$'
        return bool(re.match(plate_pattern, value.upper()))

    @staticmethod
    def _is_order_number(value: str) -> bool:
        """Проверяет номер заказа"""
        order_pattern = r'\d{8}'
        return bool(re.match(order_pattern, value.strip()))

    @staticmethod
    def _is_date(value: str) -> bool:
        """Проверяет дату"""
        date_patterns = [
            r'^\d{2}\.\d{2}\.\d{4}$',
            r'^\d{4}-\d{2}-\d{2}$',
            r'^\d{2}/\d{2}/\d{4}$'
        ]
        return any(re.match(p, value.strip()) for p in date_patterns)

    @staticmethod
    def _is_amount(value: str) -> bool:
        """Проверяет сумму"""
        amount_pattern = r'^[\d\s]+\s*₽?$|^[\d\s]+\s*руб\.?$'
        return bool(re.match(amount_pattern, value.strip()))

    def get_patterns(self) -> Dict[str, List[str]]:
        """Возвращает все извлечённые паттерны"""
        return self.patterns

    def get_search_patterns(self) -> List[str]:
        """Возвращает плоский список всех паттернов для поиска"""
        all_patterns = []
        for patterns in self.patterns.values():
            all_patterns.extend(patterns)
        return list(set(all_patterns))

    def save_to_storage(self, storage: PatternStorage, registry_info: Dict[str, Any] = None):
        """Сохраняет извлечённые паттерны в хранилище

        Args:
            storage: экземпляр PatternStorage
            registry_info: информация о реестре (путь, количество записей и т.д.)
        """
        metadata = {
            'registry_path': str(registry_info.get('path', '')) if registry_info else '',
            'record_count': registry_info.get('count', len(self.df)) if registry_info else len(self.df),
            'extracted_at': datetime.now().isoformat(),
            'columns': list(self.patterns.keys())
        }
        storage.save(self.patterns, metadata)


class RegistryBasedParser:
    """Парсер документов на основе паттернов из реестра"""

    def __init__(self, patterns_source):
        """Инициализация с паттернами

        Args:
            patterns_source: PatternExtractor, словарь паттернов или PatternStorage
        """
        if isinstance(patterns_source, PatternExtractor):
            self.patterns = patterns_source.get_search_patterns()
            self.patterns_by_type = patterns_source.get_patterns()
        elif isinstance(patterns_source, PatternStorage):
            self.patterns = patterns_source.get_flat_patterns()
            self.patterns_by_type = patterns_source.get_patterns()
        elif isinstance(patterns_source, dict):
            self.patterns = []
            for patterns_list in patterns_source.values():
                self.patterns.extend(patterns_list)
            self.patterns = list(set(self.patterns))
            self.patterns_by_type = patterns_source
        else:
            self.patterns = []
            self.patterns_by_type = {}

        print(f"🎯 RegistryBasedParser инициализирован с {len(self.patterns)} паттернами")

    def parse(self, filepath: str) -> List[Dict[str, Any]]:
        """Парсит документ, ищет совпадения с паттернами реестра

        Строгая логика для бухгалтерских документов:
        - Номер заказа = разделитель
        - Данные ищутся только после номера и до следующего номера
        """
        print(f"🔎 Анализирую файл: {Path(filepath).name}")

        # Извлекаем текст из документа
        text = self._extract_text(filepath)

        if not text:
            print("⚠️ Не удалось извлечь текст из файла")
            return []

        # Ищем номера заказов и остальные данные
        order_numbers, other_data = self._find_patterns(text)

        if not order_numbers:
            print("❌ Номера заказов не найдены")
            return []

        print(f"📋 Найдено номеров заказов: {len(order_numbers)}")

        # Группируем данные по номерам заказов (строгая логика)
        orders = self._group_into_orders(order_numbers, other_data, text)

        print(f"✅ Сформировано заказов: {len(orders)}")
        return orders

    def _extract_text(self, filepath: str) -> Optional[str]:
        """Извлекает текст из файла"""
        path = Path(filepath)
        ext = path.suffix.lower()

        try:
            if ext == '.docx':
                doc = Document(filepath)
                text = '\n'.join([p.text for p in doc.paragraphs])
                # Таблицы
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                return text

            elif ext == '.pdf' and PDF_SUPPORT:
                import pdfplumber
                with pdfplumber.open(filepath) as pdf:
                    text = ""
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        text += page_text + "\n"
                        # Таблицы
                        tables = page.extract_tables()
                        if tables:
                            for table in tables:
                                for row in table:
                                    for cell in row:
                                        if cell:
                                            text += str(cell) + " "
                    return text

            elif ext in ['.xlsx', '.xls']:
                df = pd.read_excel(filepath, dtype=str)
                return df.to_string()

            elif ext == '.csv':
                df = pd.read_csv(filepath, dtype=str)
                return df.to_string()

        except Exception as e:
            print(f"❌ Ошибка извлечения текста: {e}")

        return None

    def _find_patterns(self, text: str) -> Tuple[List[Tuple[str, int]], Dict[str, List[Tuple[str, int]]]]:
        """Ищет паттерны в тексте

        Returns:
            (список_номеров_заказов, словарь_остальных_данных)
        """
        # Сначала находим ВСЕ номера заказов (8 цифр)
        order_numbers = []
        for match in re.finditer(r'\d{8}', text):
            order_num = match.group()
            # Проверяем что это действительно номер заказа (не просто дата)
            if self._is_valid_order_number(order_num, text, match.start()):
                order_numbers.append((order_num, match.start()))

        # Затем ищем остальные паттерны
        other_data = {}
        for pattern in self.patterns:
            # Пропускаем номера заказов (уже нашли)
            if pattern.isdigit() and len(pattern) == 8:
                continue

            escaped_pattern = re.escape(pattern)
            matches = list(re.finditer(escaped_pattern, text, re.IGNORECASE))
            if matches:
                other_data[pattern] = [(m.group(), m.start()) for m in matches]

        return order_numbers, other_data

    def _is_valid_order_number(self, num: str, text: str, pos: int) -> bool:
        """Проверяет что это номер заказа, а не случайные 8 цифр"""
        # Проверяем контекст вокруг номера
        context_start = max(0, pos - 50)
        context_end = min(len(text), pos + 50)
        context = text[context_start:context_end].lower()

        # Ключевые слова которые обычно рядом с номером заказа
        order_keywords = ['заказ', 'наряд', 'заявка', 'рейс', 'номер', '№', '№№', 'req', 'order']

        # Если рядом есть ключевое слово - скорее всего это номер заказа
        if any(kw in context for kw in order_keywords):
            return True

        # Если рядом есть VIN или гос. номер - это точно заказ
        vin_near = re.search(r'[A-HJ-NPR-Z0-9]{17}', context)
        plate_near = re.search(r'[АВЕКМНОРСТУХ]\d{3}[АВЕКМНОРСТУХ]{2}\d{2,3}', context)

        if vin_near or plate_near:
            return True

        # Иначе - проверяем что не дата (дата обычно имеет формат даты вокруг)
        # Если вокруг есть года, месяцы - возможно это дата
        if re.search(r'(\d{4})|(января|февраля|марта)', context):
            return False

        return True

    def _group_into_orders(self, order_numbers: List[Tuple[str, int]], other_data: Dict[str, List[Tuple[str, int]]], text: str) -> List[Dict[str, Any]]:
        """Строгая группировка для бухгалтерских документов

        Логика:
        1. Номер заказа = разделитель
        2. Данные ищутся ТОЛЬКО после номера и ДО следующего номера
        """
        orders = []

        if not order_numbers:
            return orders

        # Сортируем номера по позиции
        order_numbers.sort(key=lambda x: x[1])

        # Обрабатываем каждый номер заказа
        for i, (order_num, order_pos) in enumerate(order_numbers):
            order = {'order_number': order_num}

            # Определяем границы для поиска данных
            # Начало = позиция номера заказа
            # Конец = позиция следующего номера (или конец текста)
            start_pos = order_pos
            end_pos = order_numbers[i + 1][1] if i + 1 < len(order_numbers) else len(text)

            # Получаем фрагмент текста для этого заказа
            order_text = text[start_pos:end_pos]

            # Ищем данные в этом фрагменте
            for pattern, matches in other_data.items():
                for value, pos in matches:
                    # Проверяем что позиция внутри границ заказа
                    if start_pos < pos < end_pos:
                        field_type = self._detect_field_type(pattern, value)
                        if field_type and field_type not in order:
                            order[field_type] = value

            orders.append(order)

        return orders

    def _detect_field_type(self, pattern: str, value: str) -> Optional[str]:
        """Определяет тип поля по паттерну/значению"""
        # Проверяем по значению
        if PatternExtractor._is_vin(value):
            return 'vin'
        elif PatternExtractor._is_plate(value):
            return 'plate'
        elif PatternExtractor._is_order_number(value):
            return 'order_number'
        elif PatternExtractor._is_date(value):
            return 'date_open'
        elif PatternExtractor._is_amount(value):
            return 'amount'

        # Проверяем по названию колонки-источника
        for col_name, patterns in self.patterns_by_type.items():
            if pattern in patterns:
                col_lower = col_name.lower()
                if 'vin' in col_lower:
                    return 'vin'
                elif 'гос' in col_lower or 'plate' in col_lower or 'номер' in col_lower:
                    return 'plate'
                elif 'заказ' in col_lower or 'order' in col_lower:
                    return 'order_number'
                elif 'дата' in col_lower or 'date' in col_lower:
                    return 'date_open'
                elif 'сумм' in col_lower or 'amount' in col_lower:
                    return 'amount'
                elif 'модель' in col_lower or 'model' in col_lower:
                    return 'model'

        return None


# ==================== ПРОВЕРКА ПО РЕЕСТРУ ====================

class AdaptiveRegistryLoader:
    """Адаптивный загрузчик реестра - определение структуры автоматически"""

    @staticmethod
    def detect_structure(registry_path: str) -> Dict[str, Any]:
        """Автоопределение структуры реестра"""
        import openpyxl

        result = {
            'sheet_name': None,
            'header_row': None,
            'columns': {},
            'order_col': None,
            'detected_patterns': []
        }

        # 1. Определяем лист
        wb = openpyxl.load_workbook(registry_path, read_only=True, data_only=True)

        # Ищем лист с ключевыми словами
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            # Проверяем первые 20 строк на наличие ключевых слов
            for row_idx in range(min(20, ws.max_row)):
                row_values = [str(cell.value or '') for cell in ws[row_idx + 1]]
                row_text = ' '.join(row_values).lower()

                # Ключевые слова для реестра заказов
                keywords = ['заказ', 'наряд', 'номер', 'дата', 'vin', 'гос', 'модель', 'обслуживание']

                if sum(1 for kw in keywords if kw in row_text) >= 3:
                    result['sheet_name'] = sheet_name
                    result['header_row'] = row_idx
                    result['detected_patterns'].append(f"Заголовок: строка {row_idx + 1}")

                    # Анализируем колонки в этой строке
                    for col_idx, cell_value in enumerate(row_values):
                        col_lower = cell_value.lower()

                        # Определяем тип колонки
                        if any(kw in col_lower for kw in ['номер заказа', 'заказ-наряд', 'заказ №', '№ заказа']):
                            result['columns']['order_number'] = col_idx
                            result['order_col'] = col_idx
                            result['detected_patterns'].append(f"№: кол {col_idx} ({cell_value})")

                        elif 'vin' in col_lower:
                            result['columns']['vin'] = col_idx
                            result['detected_patterns'].append(f"VIN: кол {col_idx} ({cell_value})")

                        elif any(kw in col_lower for kw in ['гос', 'гос. номер', 'рег', 'номер знака']):
                            result['columns']['plate'] = col_idx
                            result['detected_patterns'].append(f"Гос: кол {col_idx} ({cell_value})")

                        elif 'дата' in col_lower:
                            result['columns']['date'] = col_idx
                            result['detected_patterns'].append(f"Дата: кол {col_idx} ({cell_value})")

                        elif any(kw in col_lower for kw in ['модель', 'авто', 'автомобиль']):
                            result['columns']['model'] = col_idx
                            result['detected_patterns'].append(f"Модель: кол {col_idx} ({cell_value})")

                        elif any(kw in col_lower for kw in ['сумма', 'стоимость', 'цена', 'оплата']):
                            result['columns']['amount'] = col_idx
                            result['detected_patterns'].append(f"Сумма: кол {col_idx} ({cell_value})")

                        elif any(kw in col_lower for kw in ['вид', 'услуга', 'работ', 'обслуживание']):
                            result['columns']['service_type'] = col_idx
                            result['detected_patterns'].append(f"Услуги: кол {col_idx} ({cell_value})")

                    break  # Нашли заголовок, выходим

            if result['header_row'] is not None:
                break  # Нашли подходящий лист

        wb.close()

        # Fallback: если не удалось определить структуру
        if result['header_row'] is None:
            result['detected_patterns'].append("⚠️ Не удалось auto-определить, используем стандартную")
            result['sheet_name'] = wb.sheetnames[0] if wb.sheetnames else 'Лист1'
            result['header_row'] = 1
            result['order_col'] = 1

        return result


class RegistryChecker:
    """Проверка заказов по реестру с адаптивной загрузкой и обучением на данных"""

    def __init__(self, registry_path: str, storage_path: str = None):
        self.registry_path = Path(registry_path)
        self.df = None
        self.structure_info = None
        self.pattern_extractor = None  # Извлекатель паттернов
        self.pattern_storage = PatternStorage(storage_path)  # Хранилище паттернов
        self._load_registry(registry_path)

        # Обучаемся на реестре после загрузки
        self._train_on_registry()

    def _load_registry(self, registry_path: str):
        # Сначала пробуем автоопределение структуры
        try:
            self.structure_info = AdaptiveRegistryLoader.detect_structure(registry_path)

            sheet_name = self.structure_info['sheet_name']
            header_row = self.structure_info['header_row']

            # Загружаем с определённой структурой
            df = pd.read_excel(registry_path, sheet_name=sheet_name, header=header_row, dtype=str)
            df['excel_row'] = df.index + header_row + 2

            # Определяем колонку с номером заказа
            order_col = self.structure_info['order_col']

            if order_col is not None:
                # Используем определённую колонку
                df['order_number'] = df.iloc[:, order_col].astype(str).str.extract(r'(\d{8})')[0]
            else:
                # Fallback: ищем колонку по содержимому
                for col_idx in range(min(len(df.columns), 10)):
                    sample_values = df.iloc[:min(50, len(df)), col_idx].astype(str).str.extract(r'(\d{8})')[0]
                    if sample_values.notna().sum() > 10:  # Если нашли >10 номеров заказа
                        df['order_number'] = df.iloc[:, col_idx].astype(str).str.extract(r'(\d{8})')[0]
                        order_col = col_idx
                        break

            # Если всё ещё не нашли - пробуем все колонки
            if 'order_number' not in df.columns or df['order_number'].isna().all():
                for col in df.columns:
                    extracted = df[col].astype(str).str.extract(r'(\d{8})')[0]
                    if extracted.notna().sum() > 5:
                        df['order_number'] = extracted
                        break

            # Собираем нужные колонки по detected names
            columns_map = self.structure_info['columns']

            # Формируем итоговый dataframe с detected колонками
            final_cols = {'order_number': df['order_number'].values, 'excel_row': df['excel_row'].values}

            # Добавляем оригинальную колонку с номером заказа
            if order_col is not None and order_col < len(df.columns):
                original_order_col = df.columns[order_col]
                final_cols['original_order'] = df[original_order_col].values

            # Добавляем остальные колонки если найдены
            if 'date' in columns_map and columns_map['date'] < len(df.columns):
                final_cols['date'] = df.iloc[:, columns_map['date']].values

            if 'vin' in columns_map and columns_map['vin'] < len(df.columns):
                final_cols['vin'] = df.iloc[:, columns_map['vin']].values

            if 'plate' in columns_map and columns_map['plate'] < len(df.columns):
                final_cols['plate'] = df.iloc[:, columns_map['plate']].values

            if 'model' in columns_map and columns_map['model'] < len(df.columns):
                final_cols['model'] = df.iloc[:, columns_map['model']].values

            if 'service_type' in columns_map and columns_map['service_type'] < len(df.columns):
                final_cols['service_type'] = df.iloc[:, columns_map['service_type']].values

            if 'amount' in columns_map and columns_map['amount'] < len(df.columns):
                final_cols['amount'] = df.iloc[:, columns_map['amount']].values

            self.df = pd.DataFrame(final_cols)

            # Убедимся что есть нужные колонки (добавляем пустые если нет)
            required_cols = ['order_number', 'original_order', 'excel_row', 'date', 'vin', 'plate', 'service_type']
            for col in required_cols:
                if col not in self.df.columns:
                    if col == 'original_order':
                        self.df[col] = self.df['order_number']
                    else:
                        self.df[col] = ''

        except Exception as e:
            # Fallback к оригинальному методу если адаптивный failed
            print(f"⚠️ Адаптивная загрузка: {e}, используем стандартную")
            self._load_registry_fallback(registry_path)

        # ВАЖНО: удаляем дубликаты
        self.df = self.df.dropna(subset=['order_number']).drop_duplicates(subset=['order_number'])

    def _load_registry_fallback(self, registry_path: str):
        """Fallback метод - оригинальная логика"""
        df = pd.read_excel(registry_path, sheet_name='Лист1', header=2)
        df['excel_row'] = df.index + 4
        original_order_col = df.columns[1]
        df['order_number'] = df.iloc[:, 1].astype(str).str.extract(r'(\d{8})')[0]

        self.df = df[['order_number', original_order_col, 'excel_row',
                       df.columns[2], df.columns[3], df.columns[4], df.columns[6]]].copy()
        self.df.columns = ['order_number', 'original_order', 'excel_row',
                          'date', 'vin', 'plate', 'service_type']

    def get_structure_info(self) -> str:
        """Возвращает информацию об определённой структуре"""
        if self.structure_info:
            patterns = self.structure_info.get('detected_patterns', [])
            return " | ".join(patterns)
        return "Стандартная структура"

    def check_order(self, order: Dict[str, Any]) -> Dict[str, Any]:
        order_num = order.get('order_number')

        if not order_num:
            return {'found': False, 'reason': 'Нет номера заказа'}

        matches = self.df[self.df['order_number'] == order_num]

        if len(matches) == 0:
            return {'found': False, 'reason': 'Не найден в реестре'}

        registry_record = matches.iloc[0]

        result = {
            'found': True,
            'sheet': self.structure_info.get('sheet_name', 'Лист1') if self.structure_info else 'Лист1',
            'row': int(registry_record['excel_row']),
            'registry_order': registry_record['original_order'],
            'registry_date': str(registry_record['date'])[:10],
            'registry_vin': registry_record['vin'],
            'registry_plate': registry_record['plate'],
            'service_type': registry_record['service_type'],
            'matches': {'order_number': True}
        }

        if order.get('vin'):
            result['matches']['vin'] = order['vin'] == result['registry_vin']

        if order.get('plate'):
            result['matches']['plate'] = order['plate'] == result['registry_plate']

        if order.get('date_open'):
            try:
                order_date = datetime.strptime(order['date_open'], '%d.%m.%Y')
                reg_date = pd.to_datetime(registry_record['date'])
                result['matches']['date'] = order_date.date() == reg_date.date()
            except:
                result['matches']['date'] = None

        result['all_match'] = all(v is True for v in result['matches'].values())
        return result

    def _train_on_registry(self):
        """Обучает паттерны на основе данных реестра"""
        if self.df is None or self.df.empty:
            print("⚠️ Нет данных для обучения")
            return

        # Проверяем актуальность сохранённых паттернов
        if not self.pattern_storage.needs_update(self.registry_path, len(self.df)):
            print("✅ Используем актуальные паттерны из хранилища")
            return

        print("🧠 Обучение на реестре...")
        logger.info(f"Начинаем обучение на реестре: {self.registry_path}")
        self.pattern_extractor = PatternExtractor(self.df)

        # Сохраняем паттерны для будущего использования
        registry_info = {
            'path': str(self.registry_path),
            'count': len(self.df),
            'sheet': self.structure_info.get('sheet_name', 'Лист1') if self.structure_info else 'Лист1'
        }
        self.pattern_extractor.save_to_storage(self.pattern_storage, registry_info)
        logger.info(f"Паттерны сохранены: {len(self.df)} записей обработано")

    def get_registry_parser(self) -> Optional[RegistryBasedParser]:
        """Возвращает парсер, обученный на реестре"""
        # Приоритет: сохранённые паттерны
        if not self.pattern_storage.is_empty():
            return RegistryBasedParser(self.pattern_storage)

        # Fallback: только что извлечённые паттерны
        if self.pattern_extractor is None:
            print("⚠️ Паттерны не извлечены. Сначала загрузите реестр.")
            return None

        return RegistryBasedParser(self.pattern_extractor)

    def get_patterns_info(self) -> str:
        """Возвращает информацию об извлечённых паттернах"""
        if self.pattern_extractor is None:
            return "Паттерны не извлечены"

        patterns = self.pattern_extractor.get_patterns()
        total = sum(len(p) for p in patterns.values())

        info = f"📊 Паттернов: {total} из {len(patterns)} колонок\n"
        for col, col_patterns in patterns.items():
            info += f"  • {col}: {len(col_patterns)}\n"

        return info


# ==================== GUI ====================

class OrderCheckerApp(ctk.CTk):
    """Главное приложение v4.0 с LLM-парсером"""

    def __init__(self):
        super().__init__()

        self.title("📋 Поиск совпадений по реестру")
        self.geometry("1000x1100")  # Увеличил высоту

        # Данные
        self.registry_path = None
        self.order_files = []
        self.order_folder = None
        self.selection_mode = ctk.StringVar(value="files")
        self.checker = None
        self.processing_results = []
        self.llm_enabled = ctk.BooleanVar(value=False)
        self.llm_available = False
        self.smart_parser = None

        # Создаем интерфейс
        self._create_widgets()

        # Проверяем доступность Ollama после создания виджетов
        self._check_llm_availability()

    def _create_widgets(self):
        # Header
        header = ctk.CTkFrame(self, fg_color=COLORS['primary'], height=65)
        header.pack(fill="x")
        header.pack_propagate(False)

        ctk.CTkLabel(
            header,
            text="📋 Поиск совпадений по реестру",
            font=("Arial", 18, "bold"),
            text_color="white"
        ).pack(pady=12)

        # Контейнер
        main = ctk.CTkFrame(self, fg_color=("gray95", "gray20"))
        main.pack(fill="both", expand=True, padx=15, pady=(10, 0))

        # Секции
        self._create_registry_section(main)
        self._create_files_section(main)
        self._create_llm_section(main)
        self._create_process_section(main)
        self._create_results_section(main)

    def _create_registry_section(self, parent):
        section = ctk.CTkFrame(parent, fg_color=("white", "#2D2D2D"), corner_radius=10)
        section.pack(fill="x", pady=(0, 10))

        ctk.CTkLabel(section, text="📁 РЕЕСТР", font=("Arial", 11, "bold")).pack(anchor="w", padx=15, pady=(10, 5))

        input_frame = ctk.CTkFrame(section, fg_color="transparent")
        input_frame.pack(fill="x", padx=15, pady=(0, 10))

        self.registry_label = ctk.CTkLabel(input_frame, text="Файл не выбран", anchor="w")
        self.registry_label.pack(side="left", fill="x", expand=True, padx=10)

        ctk.CTkButton(input_frame, text="📂 Обзор", command=self._select_registry, width=70).pack(side="right", padx=10)

        # Информация о структуре реестра
        self.registry_structure_label = ctk.CTkLabel(
            section,
            text="",
            font=("Arial", 9),
            text_color="gray",
            anchor="w"
        )
        self.registry_structure_label.pack(anchor="w", padx=15, pady=(0, 10))

    def _create_files_section(self, parent):
        section = ctk.CTkFrame(parent, fg_color=("white", "#2D2D2D"), corner_radius=10)
        section.pack(fill="x", pady=(0, 10))

        ctk.CTkLabel(section, text="📄 ДОКУМЕНТЫ (DOCX, PDF, Excel, CSV)", font=("Arial", 11, "bold")).pack(anchor="w", padx=15, pady=(10, 5))

        # Селектор режима выбора
        selector_frame = ctk.CTkFrame(section, fg_color=("gray85", "gray25"))
        selector_frame.pack(fill="x", padx=15, pady=(0, 5))

        ctk.CTkRadioButton(
            selector_frame,
            text="Выбрать файлы",
            variable=self.selection_mode,
            value="files",
            command=self._on_selection_mode_change,
            font=("Arial", 10)
        ).pack(side="left", padx=10, pady=5)

        ctk.CTkRadioButton(
            selector_frame,
            text="Выбрать папку",
            variable=self.selection_mode,
            value="folder",
            command=self._on_selection_mode_change,
            font=("Arial", 10)
        ).pack(side="left", padx=10, pady=5)

        # Зона файлов
        files_frame = ctk.CTkFrame(section, fg_color="transparent")
        files_frame.pack(fill="x", padx=15, pady=(0, 5))

        self.files_label = ctk.CTkLabel(files_frame, text="Не выбрано", anchor="w")
        self.files_label.pack(side="left", fill="x", expand=True, padx=10)

        self.select_btn = ctk.CTkButton(files_frame, text="📂 Обзор", command=self._select_files, width=70)
        self.select_btn.pack(side="right", padx=10)

        # Информация о выбранных файлах
        self.files_info_frame = ctk.CTkFrame(section, fg_color="transparent")
        self.files_info_frame.pack(fill="x", padx=15, pady=(0, 10))

        self.files_info_label = ctk.CTkLabel(
            self.files_info_frame,
            text="",
            font=("Arial", 9),
            text_color="gray"
        )
        self.files_info_label.pack(anchor="w", padx=10, pady=5)

    def _create_llm_section(self, parent):
        """Секция настроек LLM"""
        section = ctk.CTkFrame(parent, fg_color=("white", "#2D2D2D"), corner_radius=10)
        section.pack(fill="x", pady=(0, 10))

        # Заголовок
        header = ctk.CTkFrame(section, fg_color="transparent")
        header.pack(fill="x", padx=15, pady=(10, 5))

        ctk.CTkLabel(header, text="🤖 УМНЫЙ ПАРСЕР (LLM)", font=("Arial", 11, "bold")).pack(side="left")

        self.llm_status_label = ctk.CTkLabel(
            header,
            text="",
            font=("Arial", 9),
            text_color="gray"
        )
        self.llm_status_label.pack(side="right")

        # Тogle для LLM
        toggle_frame = ctk.CTkFrame(section, fg_color="transparent")
        toggle_frame.pack(fill="x", padx=15, pady=(0, 10))

        ctk.CTkSwitch(
            toggle_frame,
            text="Использовать Ollama для сложных документов",
            variable=self.llm_enabled,
            command=self._on_llm_toggle,
            font=("Arial", 10)
        ).pack(side="left", padx=10)

        # Инфо
        info_frame = ctk.CTkFrame(section, fg_color=("gray85", "gray25"))
        info_frame.pack(fill="x", padx=15, pady=(0, 10))

        ctk.CTkLabel(
            info_frame,
            text="💡 LLM автоматически используется для документов с нестандартной структурой",
            font=("Arial", 9),
            text_color="gray"
        ).pack(anchor="w", padx=10, pady=5)

    def _check_llm_availability(self):
        """Проверяет доступность Ollama"""
        try:
            client = OllamaClient()
            if client.is_available():
                self.llm_available = True
                self.llm_status_label.configure(text="✅ Ollama готов", text_color=COLORS['success'])
            else:
                self.llm_status_label.configure(text="⚠️ Ollama не найден", text_color=COLORS['warning'])
        except Exception:
            self.llm_available = False
            self.llm_status_label.configure(text="❌ Недоступно", text_color=COLORS['error'])

    def _on_llm_toggle(self):
        """Обработка переключения LLM"""
        if self.llm_enabled.get():
            if not self.llm_available:
                messagebox.showwarning(
                    "Ollama недоступен",
                    "Ollama не запущен или не установлен.\n\n"
                    "1. Установите Ollama: https://ollama.ai\n"
                    "2. Запустите: ollama serve\n"
                    "3. Скачайте модель: ollama pull llama3.2"
                )
                self.llm_enabled.set(False)
                self.llm_status_label.configure(text="❌ Недоступно", text_color=COLORS['error'])
            else:
                # Создаём SmartParser с LLM
                self.smart_parser = SmartParser({'llm_enabled': True})
                self.llm_status_label.configure(text="✅ LLM включён", text_color=COLORS['success'])
        else:
            self.smart_parser = None
            self.llm_status_label.configure(text="✅ Ollama готов", text_color=COLORS['success'])

    def _create_process_section(self, parent):
        section = ctk.CTkFrame(parent, fg_color=("white", "#2D2D2D"), corner_radius=10)
        section.pack(fill="x", pady=(0, 10))

        self.process_btn = ctk.CTkButton(
            section,
            text="🚀 ВЫПОЛНИТЬ СВЕРКУ",
            command=self._process,
            state="disabled",
            height=45,
            font=("Arial", 13, "bold"),
            fg_color=COLORS['success']
        )
        self.process_btn.pack(fill="x", padx=15, pady=10)

        # Прогресс
        self.progress_label = ctk.CTkLabel(section, text="", text_color="gray")
        self.progress_label.pack(anchor="w", padx=15, pady=(0, 5))

        self.progress_bar = ctk.CTkProgressBar(section, height=10)
        self.progress_bar.pack(fill="x", padx=15, pady=(0, 10))
        self.progress_bar.set(0)

    def _create_results_section(self, parent):
        section = ctk.CTkFrame(parent, fg_color=("white", "#2D2D2D"), corner_radius=10)
        section.pack(fill="both", expand=True, padx=(15, 10), pady=(0, 15))

        # Header результатов
        results_header = ctk.CTkFrame(section, fg_color="transparent")
        results_header.pack(fill="x", padx=15, pady=(10, 5))

        ctk.CTkLabel(results_header, text="📋 РЕЗУЛЬТАТЫ", font=("Arial", 12, "bold")).pack(side="left")

        self.stats_label = ctk.CTkLabel(results_header, text="", font=("Arial", 10))
        self.stats_label.pack(side="right")

        # Кнопки экспорта (НАД результатами)
        actions = ctk.CTkFrame(section, fg_color="transparent")
        actions.pack(fill="x", padx=15, pady=(0, 10))

        self.export_btn = ctk.CTkButton(
            actions,
            text="💾 Экспорт",
            command=self._export_dialog,
            state="disabled",
            height=40,
            font=("Arial", 11, "bold"),
            fg_color=COLORS['primary']
        )
        self.export_btn.pack(side="left", fill="x", expand=True, padx=(0, 5))

        ctk.CTkButton(actions, text="🗑️ Очистить", command=self._clear, height=40, font=("Arial", 11)).pack(side="left", fill="x", expand=True, padx=(5, 0))

        # Таблица результатов (с фиксированной высотой для скролла)
        self.results_table = ctk.CTkScrollableFrame(section, fg_color="transparent", height=300)
        self.results_table.pack(fill="both", expand=True, padx=10, pady=(0, 15))

    def _select_registry(self):
        filepath = filedialog.askopenfilename(
            filetypes=[("Excel", "*.xlsx *.xls"), ("Все", "*.*")],
            title="Выберите реестр"
        )
        if filepath:
            self.registry_path = filepath
            self.registry_label.configure(text=Path(filepath).name)
            try:
                self.checker = RegistryChecker(filepath)
                total = len(self.checker.df)
                structure_info = self.checker.get_structure_info()
                self.registry_structure_label.configure(text=f"🔍 Структура: {structure_info}")
                self._check_ready()
                messagebox.showinfo("Реестр", f"Загружено заказов: {total}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка:\n{e}")
                self.registry_structure_label.configure(text="")

    def _on_selection_mode_change(self):
        """Обработка изменения режима выбора"""
        mode = self.selection_mode.get()
        if mode == "files":
            self.select_btn.configure(text="📂 Обзор")
            self.files_label.configure(text="Файлы не выбраны" if not self.order_files else f"Выбрано: {len(self.order_files)}")
        else:
            self.select_btn.configure(text="📁 Выбрать папку")
            self.files_label.configure(text="Папка не выбрана" if not self.order_folder else f"Папка: {Path(self.order_folder).name}")

    def _select_files(self):
        """Выбор файлов или папки с документами"""
        mode = self.selection_mode.get()

        if mode == "files":
            # Выбор файлов
            filepaths = filedialog.askopenfilenames(
                filetypes=[
                    ("Документы", "*.docx *.pdf *.xlsx *.xls *.csv"),
                    ("Word", "*.docx"),
                    ("PDF", "*.pdf"),
                    ("Excel", "*.xlsx *.xls"),
                    ("CSV", "*.csv"),
                    ("Все", "*.*")
                ],
                title="Выберите файлы (можно несколько)"
            )

            if filepaths:
                self.order_files = list(filepaths)
                self.order_folder = None
                self._update_files_info()

        else:
            # Выбор папки
            folderpath = filedialog.askdirectory(title="Выберите папку с документами")

            if folderpath:
                self.order_folder = folderpath
                # Находим все поддерживаемые файлы в папке
                self.order_files = []
                folder_path = Path(folderpath)

                # Ищем DOCX, PDF, Excel, CSV файлы
                for ext in ['*.docx', '*.pdf', '*.xlsx', '*.xls', '*.csv']:
                    self.order_files.extend(str(f) for f in folder_path.glob(ext))
                    # Также ищем в подпапках
                    self.order_files.extend(str(f) for f in folder_path.rglob(ext))

                self._update_files_info()

        self._check_ready()

    def _update_files_info(self):
        """Обновляет информацию о выбранных файлах"""
        mode = self.selection_mode.get()

        if mode == "files":
            count = len(self.order_files)
            self.files_label.configure(text=f"Выбрано: {count}")

            # Подсчитаем по типам
            docx_count = sum(1 for f in self.order_files if f.endswith('.docx'))
            pdf_count = sum(1 for f in self.order_files if f.endswith('.pdf'))
            excel_count = sum(1 for f in self.order_files if f.endswith(('.xlsx', '.xls')))
            csv_count = sum(1 for f in self.order_files if f.endswith('.csv'))

            info_text = f"DOCX: {docx_count} | PDF: {pdf_count} | Excel: {excel_count} | CSV: {csv_count}"
            self.files_info_label.configure(text=info_text)
        else:
            self.files_label.configure(text=f"Папка: {Path(self.order_folder).name}")
            count = len(self.order_files)
            docx_count = sum(1 for f in self.order_files if f.endswith('.docx'))
            pdf_count = sum(1 for f in self.order_files if f.endswith('.pdf'))
            excel_count = sum(1 for f in self.order_files if f.endswith(('.xlsx', '.xls')))
            csv_count = sum(1 for f in self.order_files if f.endswith('.csv'))
            info_text = f"Найдено: {count} (DOCX: {docx_count} | PDF: {pdf_count} | Excel: {excel_count} | CSV: {csv_count})"
            self.files_info_label.configure(text=info_text)

    def _check_ready(self):
        if self.registry_path and self.order_files:
            self.process_btn.configure(state="normal")

    def _process(self):
        if not self.registry_path or not self.order_files:
            messagebox.showwarning("Внимание", "Выберите реестр и файлы")
            return

        self.process_btn.configure(state="disabled")
        self._clear_results()

        thread = threading.Thread(target=self._process_thread, daemon=True)
        thread.start()

    def _process_thread(self):
        try:
            if not self.checker:
                self.checker = RegistryChecker(self.registry_path)

            # Создаём SmartParser если включен LLM
            parser_config = {'llm_enabled': self.llm_enabled.get()}
            if parser_config['llm_enabled'] and not self.smart_parser:
                self.smart_parser = SmartParser(parser_config)

            # Получаем парсер, обученный на реестре
            registry_parser = self.checker.get_registry_parser()
            if registry_parser:
                print("✅ Используется парсер, обученный на реестре")

            self.progress_label.configure(text="Обработка...")
            self.progress_bar.set(0.05)

            total = len(self.order_files)
            results = []
            stats = {'total': 0, 'found': 0, 'matched': 0, 'not_found': 0, 'llm': 0, 'registry': 0, 'new': 0}

            for i, filepath in enumerate(self.order_files):
                try:
                    # Показываем что обрабатываем
                    filename = Path(filepath).name
                    self.progress_label.configure(text=f"Обработка: {filename}")

                    # === НОВАЯ ЛОГИКА: Универсальный парсинг с приоритетом на НОВЫЕ данные ===
                    orders = []

                    # 1. ПРИОРИТЕТ: LLM если доступен (универсальный парсинг)
                    if self.smart_parser and self.llm_enabled.get() and self.llm_available:
                        orders = self.smart_parser.parse(filepath)
                        if orders:
                            stats['llm'] += len(orders)
                            parser_type = "🤖"
                            print(f"  🤖 Найдено через LLM: {len(orders)}")

                    # 2. Универсальный парсер (извлекает ВСЕ данные по структуре)
                    if not orders:
                        orders = UniversalParser.parse(filepath)
                        if orders:
                            parser_type = "📋"
                            print(f"  📋 Найдено через универсальный парсер: {len(orders)}")

                    # 3. Проверяем по реестру (дополнительно найденные данные)
                    # Каждому найденному заказу добавляем метку если он из реестра
                    for order in orders:
                        if registry_parser:
                            # Проверяем есть ли этот заказ в реестре
                            check_result = self.checker.check_order(order)
                            if check_result.get('found'):
                                order['from_registry'] = True
                                order['registry_info'] = check_result
                            else:
                                order['from_registry'] = False
                                order['registry_info'] = check_result
                        else:
                            order['from_registry'] = False

                    for order in orders:
                        check_result = self.checker.check_order(order)

                        stats['total'] += 1
                        if check_result['found']:
                            stats['found'] += 1
                            if check_result.get('all_match'):
                                stats['matched'] += 1
                        else:
                            stats['not_found'] += 1
                            # Подсчитываем НОВЫЕ заказы (которых нет в реестре)
                            if order.get('order_number'):
                                stats['new'] = stats.get('new', 0) + 1

                        self._add_result(Path(filepath).name, order, check_result)

                    progress = (i + 1) / total
                    self.progress_bar.set(0.1 + progress * 0.85)
                    self.progress_label.configure(text=f"Обработано: {i+1}/{total}")

                except Exception as e:
                    print(f"Ошибка {filepath}: {e}")

            # Подсчитываем статистику LLM
            llm_count = sum(1 for r in self.processing_results if r.get('is_llm', False))
            new_count = stats.get('new', 0)

            # Итоговая статистика с НОВЫМИ заказами
            llm_text = f" | 🤖 LLM: {llm_count}" if llm_count > 0 else ""
            new_text = f" | 🆕 НОВЫЕ: {new_count}" if new_count > 0 else ""
            stats_text = f"Всего: {stats['total']} | ✅: {stats['found']} | 🆕: {new_count}{llm_text}{new_text}"
            self.stats_label.configure(text=stats_text)
            self.progress_label.configure(text="Готово!")
            self.progress_bar.set(1)
            self.process_btn.configure(state="normal")
            self.export_btn.configure(state="normal")
            print(f"✅ Кнопка экспорта включена. Результатов: {len(self.processing_results)}")

        except Exception as e:
            import traceback
            print(f"❌ Ошибка обработки: {e}")
            traceback.print_exc()
            messagebox.showerror("Ошибка", f"Ошибка:\n\n{e}")
            self.progress_label.configure(text="Ошибка!")
            self.process_btn.configure(state="normal")
            # Включаем экспорт если есть результаты
            if hasattr(self, 'processing_results') and self.processing_results:
                self.export_btn.configure(state="normal")
                print(f"⚠️ Кнопка экспорта включена после ошибки. Результатов: {len(self.processing_results)}")
            else:
                print(f"❌ Нет результатов для экспорта")

    def _add_result(self, filename, order, check_result):
        """Добавляет результат в таблицу (thread-safe)"""
        # Проверяем - использовался ли LLM
        source_info = order.get('source_info', {})
        is_llm = any('LLM' in str(v) for v in source_info.values() if v)
        is_new = order.get('from_registry') == False
        llm_mark = ' 🤖' if is_llm else ''
        new_mark = ' 🆕' if is_new else ''

        try:
            frame = ctk.CTkFrame(self.results_table, fg_color=("gray95", "gray25"), corner_radius=8)
            frame.pack(fill="x", padx=5, pady=3)

            if check_result['found']:
                if check_result.get('all_match'):
                    status = f"✅ {check_result['sheet']}, строка {check_result['row']}"
                    color = COLORS['success']
                else:
                    status = f"⚠️ Расхождение: {check_result['sheet']}, строка {check_result['row']}"
                    color = COLORS['warning']
            else:
                # Показываем как НОВЫЙ заказ если есть номер
                if order.get('order_number'):
                    status = f"🆕 НОВЫЙ (нет в реестре)"
                    color = COLORS['info']  # Голубой для новых
                else:
                    status = "❌ Не удалось распознать"
                    color = COLORS['error']

            ctk.CTkLabel(
                frame,
                text=f"{filename} | №{order.get('order_number', '???')} | {status}{llm_mark}{new_mark}",
                font=("Arial", 10),
                text_color=color,
                anchor="w"
            ).pack(anchor="w", padx=10, pady=6)

        except Exception as e:
            # Игнорируем ошибки GUI при многопоточном обновлении
            pass

        # Всегда добавляем в данные
        self.processing_results.append({
            'filename': filename,
            'order': order,
            'check_result': check_result,
            'is_llm': is_llm,
            'is_new': is_new
        })

    def _clear_results(self):
        for widget in self.results_table.winfo_children():
            widget.destroy()
        self.processing_results = []

    def _clear(self):
        self._clear_results()
        self.progress_label.configure(text="")
        self.progress_bar.set(0)
        self.stats_label.configure(text="")
        self.files_label.configure(text="Не выбрано")
        self.files_info_label.configure(text="")
        self.order_folder = None
        self.registry_structure_label.configure(text="")
        self.smart_parser = None
        self.export_btn.configure(state="disabled")

    def _export_dialog(self):
        """Диалог выбора формата экспорта"""
        if not hasattr(self, 'processing_results') or not self.processing_results:
            messagebox.showwarning("Внимание", "Нет данных для экспорта")
            return

        # Создаем диалог
        export_dialog = ctk.CTkToplevel(self)
        export_dialog.title("💾 Экспорт отчета")
        export_dialog.geometry("400x320")
        export_dialog.transient(self)
        export_dialog.grab_set()

        # Центрируем
        export_dialog.update()
        x = (export_dialog.winfo_screenwidth() // 2) - 200
        y = (export_dialog.winfo_screenheight() // 2) - 160
        export_dialog.geometry(f"400x320+{x}+{y}")

        # Сохраняем ссылку на диалог для закрытия после экспорта
        self._export_dialog_window = export_dialog

        # Заголовок
        ctk.CTkLabel(
            export_dialog,
            text="💾 Выберите формат экспорта",
            font=("Arial", 14, "bold")
        ).pack(pady=15)

        # Кнопки форматов
        ctk.CTkButton(
            export_dialog,
            text="📊 Excel (.xlsx)",
            command=lambda: self._export_excel(export_dialog),
            width=200,
            height=40,
            font=("Arial", 11)
        ).pack(fill="x", padx=40, pady=5)

        ctk.CTkButton(
            export_dialog,
            text="📋 CSV (.csv)",
            command=lambda: self._export_csv(export_dialog),
            width=200,
            height=40,
            font=("Arial", 11)
        ).pack(fill="x", padx=40, pady=5)

        ctk.CTkButton(
            export_dialog,
            text="📄 JSON (.json)",
            command=lambda: self._export_json(export_dialog),
            width=200,
            height=40,
            font=("Arial", 11)
        ).pack(fill="x", padx=40, pady=5)

        # Кнопка отмены
        ctk.CTkButton(
            export_dialog,
            text="Отмена",
            command=export_dialog.destroy,
            width=120,
            height=35,
            fg_color="gray"
        ).pack(pady=(15, 10))

    def _export_excel(self, dialog):
        """Экспорт в Excel"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            initialfile=f"отчет_{timestamp}.xlsx",
            filetypes=[("Excel", "*.xlsx")],
            title="Экспорт в Excel",
            parent=self
        )

        if filepath:
            self._export_to_excel(filepath)
            dialog.destroy()
            messagebox.showinfo("Готово", "Отчет сохранен!")

    def _export_csv(self, dialog):
        """Экспорт в CSV"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = filedialog.asksaveasfilename(
            defaultextension=".csv",
            initialfile=f"отчет_{timestamp}.csv",
            filetypes=[("CSV", "*.csv")],
            title="Экспорт в CSV",
            parent=self
        )

        if filepath:
            self._export_to_csv(filepath)
            dialog.destroy()
            messagebox.showinfo("Готово", "Отчет сохранен!")

    def _export_json(self, dialog):
        """Экспорт в JSON"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = filedialog.asksaveasfilename(
            defaultextension=".json",
            initialfile=f"отчет_{timestamp}.json",
            filetypes=[("JSON", "*.json")],
            title="Экспорт в JSON",
            parent=self
        )

        if filepath:
            self._export_to_json(filepath)
            dialog.destroy()
            messagebox.showinfo("Готово", "Отчет сохранен!")

    def _export_to_excel(self, filepath: str):
        """Экспорт данных в Excel с подсветкой строк и детальной сверкой"""
        from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        results = self.processing_results

        # Цвета для подсветки
        COLOR_GREEN = 'C6E0B4'      # Зелёный - полное совпадение
        COLOR_YELLOW = 'FFE699'     # Жёлтый - частичное совпадение
        COLOR_RED = 'F4B183'        # Красный - не найден
        COLOR_GRAY = 'D9D9D9'       # Серый - заголовок
        COLOR_HEADER = '4472C4'     # Синий - основная шапка

        with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
            # ===== ЛИСТ 1: СВОДКА =====
            summary_data = []
            by_file = {}
            for result in results:
                filename = result['filename']
                if filename not in by_file:
                    by_file[filename] = {'total': 0, 'matched': 0, 'partial': 0, 'not_found': 0}
                by_file[filename]['total'] += 1
                check_result = result['check_result']
                if not check_result['found']:
                    by_file[filename]['not_found'] += 1
                elif check_result.get('all_match'):
                    by_file[filename]['matched'] += 1
                else:
                    by_file[filename]['partial'] += 1

            for filename, stats in by_file.items():
                summary_data.append({
                    'Файл': filename,
                    'Всего заказов': stats['total'],
                    '✅ Полное совпадение': stats['matched'],
                    '⚠️ Расхождения': stats['partial'],
                    '❌ Не найдено': stats['not_found']
                })

            total = len(results)
            matched = sum(1 for r in results if r['check_result'].get('all_match'))
            partial = sum(1 for r in results if r['check_result']['found'] and not r['check_result'].get('all_match'))
            not_found = sum(1 for r in results if not r['check_result']['found'])

            summary_data.append({
                'Файл': '📊 ИТОГО',
                'Всего заказов': total,
                '✅ Полное совпадение': matched,
                '⚠️ Расхождения': partial,
                '❌ Не найдено': not_found
            })

            df_summary = pd.DataFrame(summary_data)
            df_summary.to_excel(writer, sheet_name='Сводка', index=False)

            # Форматируем сводку
            ws_summary = writer.sheets['Сводка']
            self._format_summary_sheet(ws_summary, COLOR_HEADER, COLOR_GREEN, COLOR_YELLOW, COLOR_RED)

            # ===== ЛИСТ 2: ВСЕ ЗАКАЗЫ с детальной сверкой =====
            detail_data = []
            row_colors = []
            for result in results:
                check = result['check_result']
                order = result['order']

                # Определяем статус по полям
                order_match = check.get('matches', {}).get('order_number', True)
                vin_match = check.get('matches', {}).get('vin')
                plate_match = check.get('matches', {}).get('plate')
                date_match = check.get('matches', {}).get('date')
                amount_match = None  # TODO: добавить сравнение сумм

                # Формируем статусы для отображения
                if check['found']:
                    if check.get('all_match'):
                        overall_status = '✅ ПОЛНОЕ СОВПАДЕНИЕ'
                        row_color = 'green'
                    else:
                        overall_status = '⚠️ РАСХОЖДЕНИЕ'
                        row_color = 'yellow'
                else:
                    overall_status = '❌ НЕ НАЙДЕН В РЕЕСТРЕ'
                    row_color = 'red'

                # Получаем информацию об источнике данных
                source_info = order.get('source_info', {})

                # Формируем строку источника
                source_parts = []
                if source_info.get('order_number'):
                    source_parts.append(f"№: {source_info['order_number']}")
                if source_info.get('vin'):
                    source_parts.append(f"VIN: {source_info['vin']}")
                if source_info.get('plate'):
                    source_parts.append(f"Гос: {source_info['plate']}")
                if source_info.get('model'):
                    source_parts.append(f"Модель: {source_info['model']}")
                if source_info.get('amount'):
                    source_parts.append(f"Сумма: {source_info['amount']}")

                source_str = ' | '.join(source_parts) if source_parts else 'Не определено'

                # Данные из документа
                row = {
                    '📢 Статус': overall_status,
                    '📄 Файл документа': result['filename'],
                    '📍 Источник данных': source_str,
                    '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━': '',  # Разделитель
                    '📋 № Заказа (документ)': order['order_number'],
                    '📅 Дата открытия (документ)': order.get('date_open', ''),
                    '📅 Дата закрытия (документ)': order.get('date_close', ''),
                    '🚗 VIN код (документ)': order.get('vin', ''),
                    '🔢 Гос. номер (документ)': order.get('plate', ''),
                    '📦 Модель (документ)': order.get('model', ''),
                    '💰 Сумма (документ)': order.get('amount', ''),
                }

                if check['found']:
                    # Данные из реестра
                    row['━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'] = ''
                    row['📊 Строка в реестре'] = check.get('row', '')
                    row['📋 № Заказа (реестр)'] = check.get('registry_order', '')
                    row['📅 Дата (реестр)'] = check.get('registry_date', '')
                    row['🚗 VIN код (реестр)'] = check.get('registry_vin', '')
                    row['🔢 Гос. номер (реестр)'] = check.get('registry_plate', '')
                    row['📦 Вид обслуживания (реестр)'] = check.get('service_type', '')
                    row['💰 Сумма (реестр)'] = ''  # TODO: добавить сумму из реестра
                else:
                    row['━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'] = ''
                    row['📊 Строка в реестре'] = '❌ НЕ НАЙДЕН'
                    row['📋 № Заказа (реестр)'] = '—'
                    row['📅 Дата (реестр)'] = '—'
                    row['🚗 VIN код (реестр)'] = '—'
                    row['🔢 Гос. номер (реестр)'] = '—'
                    row['📦 Вид обслуживания (реестр)'] = '—'
                    row['💰 Сумма (реестр)'] = '—'

                # Статусы по полям (для быстрой визуальной сверки)
                row['━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'] = ''
                row['🔢 Сверка по №'] = '✅ СОВПАЛ' if order_match else '❌ НЕ СОВПАЛ'
                row['🚗 Сверка по VIN'] = '✅ СОВПАЛ' if vin_match is True else ('❌ РАСХОЖДЕНИЕ' if vin_match is False else '— НЕТ ДАННЫХ')
                row['🔢 Сверка по гос. номеру'] = '✅ СОВПАЛ' if plate_match is True else ('❌ РАСХОЖДЕНИЕ' if plate_match is False else '— НЕТ ДАННЫХ')
                row['📅 Сверка по дате'] = '✅ СОВПАЛА' if date_match is True else ('❌ РАСХОЖДЕНИЕ' if date_match is False else '— НЕТ ДАННЫХ')
                row['💰 Сверка по сумме'] = '✅ СОВПАЛА' if amount_match is True else ('❌ РАСХОЖДЕНИЕ' if amount_match is False else '— НЕТ ДАННЫХ')

                detail_data.append(row)
                row_colors.append(row_color)

            df_detail = pd.DataFrame(detail_data)
            df_detail.to_excel(writer, sheet_name='Все заказы', index=False)

            # Форматируем детальный лист с подсветкой строк
            ws_detail = writer.sheets['Все заказы']
            self._format_detail_sheet(ws_detail, row_colors, COLOR_HEADER, COLOR_GREEN, COLOR_YELLOW, COLOR_RED)

            # ===== ЛИСТ 3: ПРОБЛЕМНЫЕ =====
            issues_data = []
            issues_colors = []
            for idx, row in enumerate(detail_data):
                if row['📢 Статус'] in ['⚠️ РАСХОЖДЕНИЕ', '❌ НЕ НАЙДЕН']:
                    issues_data.append(row)
                    issues_colors.append(row_colors[idx])

            if issues_data:
                df_issues = pd.DataFrame(issues_data)
                df_issues.to_excel(writer, sheet_name='Проблемные', index=False)

                # Форматируем проблемный лист
                ws_issues = writer.sheets['Проблемные']
                self._format_detail_sheet(ws_issues, issues_colors, COLOR_HEADER, COLOR_GREEN, COLOR_YELLOW, COLOR_RED)

    def _format_summary_sheet(self, ws, header_color, green, yellow, red):
        """Форматирование листа сводки"""
        from openpyxl.styles import Font, PatternFill, Alignment
        from openpyxl.utils import get_column_letter

        # Заголовок
        for cell in ws[1]:
            cell.fill = PatternFill(start_color=header_color, end_color=header_color, fill_type='solid')
            cell.font = Font(bold=True, color='FFFFFFFF')  # Белый в ARGB
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # Итого строка
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
            if row[0].value == '📊 ИТОГО':
                for cell in row:
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color='FFE7E6E6', end_color='FFE7E6E6', fill_type='solid')

        # Автоширина
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width

    def _format_detail_sheet(self, ws, row_colors, header_color, green, yellow, red):
        """Форматирование детального листа с подсветкой строк"""
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        # Заголовок - синий
        for cell in ws[1]:
            cell.fill = PatternFill(start_color=header_color, end_color=header_color, fill_type='solid')
            cell.font = Font(bold=True, color='FFFFFFFF', size=11)  # Белый в ARGB
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

        # Подсветка строк по статусу
        for idx, row_color in enumerate(row_colors, start=2):
            fill_color = green if row_color == 'green' else (yellow if row_color == 'yellow' else red)

            for cell in ws[idx]:
                cell.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type='solid')
                cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                cell.border = Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )

        # Автоширина колонок
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 30)
            ws.column_dimensions[column_letter].width = adjusted_width

        # Выравнивание заголовка по центру
        ws.row_dimensions[1].height = 30

    def _export_to_csv(self, filepath: str):
        """Экспорт в CSV"""
        data = []
        for result in self.processing_results:
            check = result['check_result']
            order = result['order']

            row = {
                'Файл': result['filename'],
                '№ Заказа': order['order_number'],
                'Статус': 'НАЙДЕН' if check['found'] else 'НЕ НАЙДЕН'
            }

            if check['found']:
                row['Лист реестра'] = check.get('sheet', '')
                row['Строка'] = check.get('row', '')
                row['Дата реестра'] = check.get('registry_date', '')
                row['VIN реестр'] = check.get('registry_vin', '')
                row['Гос. номер реестр'] = check.get('registry_plate', '')
                row['Модель реестра'] = check.get('service_type', '')

            data.append(row)

        df = pd.DataFrame(data)
        df.to_csv(filepath, index=False, encoding='utf-8-sig')

    def _export_to_json(self, filepath: str):
        """Экспорт в JSON"""
        from collections import defaultdict

        # Группировка по файлам
        by_file = defaultdict(list)
        for result in self.processing_results:
            filename = result['filename']
            check = result['check_result']
            order = result['order']

            by_file[filename].append({
                'order_number': order['order_number'],
                'found': check['found'],
                'registry_info': {
                    'sheet': check.get('sheet', ''),
                    'row': check.get('row', ''),
                    'date': check.get('registry_date', ''),
                    'vin': check.get('registry_vin', ''),
                    'plate': check.get('registry_plate', ''),
                } if check['found'] else None
            })

        # Статистика
        total = sum(len(orders) for orders in by_file.values())
        found = sum(sum(1 for o in orders if o['found']) for orders in by_file.values())

        # Создаем структуру
        report = {
            'timestamp': datetime.now().isoformat(),
            'statistics': {
                'total_files': len(by_file),
                'total_orders': total,
                'found': found,
                'not_found': total - found
            },
            'results': dict(by_file)
        }

        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)


def main():
    app = OrderCheckerApp()
    app.mainloop()


if __name__ == "__main__":
    main()
